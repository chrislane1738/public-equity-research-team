from pathlib import Path

from docx import Document

from backend.tools.docx_writer import write_memo


def test_write_memo_produces_valid_docx_with_sections(tmp_path):
    out_path = tmp_path / "memo.docx"
    sections = [
        ("Executive Summary", "We rate NVDA Buy with $X PT."),
        ("Investment Thesis", "Three reasons we like the name..."),
        ("Risks", "Top risk: AI capex pullback."),
    ]
    write_memo(out_path, title="NVDA — Initiation", sections=sections)

    assert out_path.exists()
    doc = Document(out_path)
    paragraphs = [p.text for p in doc.paragraphs]
    assert "NVDA — Initiation" in paragraphs
    assert "Executive Summary" in paragraphs
    assert "We rate NVDA Buy with $X PT." in paragraphs
    assert "Risks" in paragraphs


def test_write_memo_handles_markdown_paragraphs(tmp_path):
    out_path = tmp_path / "memo.docx"
    body = "First paragraph.\n\nSecond paragraph with **bold**.\n\nThird."
    write_memo(out_path, title="Test", sections=[("Section", body)])

    doc = Document(out_path)
    paragraphs = [p.text for p in doc.paragraphs]
    assert "First paragraph." in paragraphs
    assert "Third." in paragraphs
