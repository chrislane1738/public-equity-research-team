"""python-docx wrapper for writing memo documents.

Plan A: minimal — title + sections of (heading, body) tuples. Markdown bold/italic
markers are stripped (Plan B will add proper inline formatting).
"""
import re
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.shared import Pt


MARKDOWN_INLINE = re.compile(r"\*\*(.*?)\*\*|\*(.*?)\*|_(.*?)_")


def _strip_markdown(text: str) -> str:
    return MARKDOWN_INLINE.sub(lambda m: m.group(1) or m.group(2) or m.group(3), text)


def write_memo(
    path: Path,
    title: str,
    sections: Sequence[tuple[str, str]],
) -> None:
    """Write a docx file with a title page and section headings + body paragraphs."""
    doc = Document()

    title_p = doc.add_paragraph()
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    for heading, body in sections:
        h = doc.add_paragraph()
        h_run = h.add_run(heading)
        h_run.bold = True
        h_run.font.size = Pt(14)

        for para in body.split("\n\n"):
            doc.add_paragraph(_strip_markdown(para.strip()))

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
